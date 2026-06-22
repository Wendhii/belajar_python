import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

def create_element(name):
    return OxmlElement(name)

def set_cell_background(cell, color_hex):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{color_hex}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=100, bottom=100, left=150, right=150):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def set_table_borders(table, color="D3D3D3"):
    tblPr = table._tbl.tblPr
    borders = parse_xml(f'''
        <w:tblBorders {nsdecls("w")}>
            <w:top w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:bottom w:val="single" w:sz="6" w:space="0" w:color="{color}"/>
            <w:left w:val="none"/>
            <w:right w:val="none"/>
            <w:insideH w:val="single" w:sz="4" w:space="0" w:color="{color}"/>
            <w:insideV w:val="none"/>
        </w:tblBorders>
    ''')
    tblPr.append(borders)

# Initialize document
doc = docx.Document()

# Base geometry
section = doc.sections[0]
section.page_width = Inches(8.27)  # A4
section.page_height = Inches(11.69)
section.top_margin = Inches(1.0)
section.bottom_margin = Inches(1.0)
section.left_margin = Inches(1.0)
section.right_margin = Inches(1.0)

# Colors
COLOR_PRIMARY = RGBColor(26, 54, 93)     # #1A365D - Dark Navy
COLOR_SECONDARY = RGBColor(43, 108, 176) # #2B6CB0 - Steel Blue
COLOR_TEXT = RGBColor(45, 55, 72)        # #2D3748 - Charcoal
HEX_BG_LIGHT = "F7FAFC"
HEX_BORDER = "E2E8F0"

# Custom Typography Setup
def apply_run_font(run, name="Arial", size_pt=11, bold=False, italic=False, color=COLOR_TEXT):
    run.font.name = name
    run.font.size = Pt(size_pt)
    run.bold = bold
    run.italic = italic
    run.font.color.rgb = color

# Title Paragraph
p_title = doc.add_paragraph()
p_title.paragraph_format.space_before = Pt(0)
p_title.paragraph_format.space_after = Pt(4)
p_title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_title = p_title.add_run("MATERI KULIAH: PIPELINE INSTRUCTION HAZARD")
apply_run_font(run_title, size_pt=18, bold=True, color=COLOR_PRIMARY)

# Subtitle Paragraph
p_sub = doc.add_paragraph()
p_sub.paragraph_format.space_before = Pt(0)
p_sub.paragraph_format.space_after = Pt(24)
p_sub.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
run_sub = p_sub.add_run("Mata Kuliah: Arsitektur dan Organisasi Komputer")
apply_run_font(run_sub, size_pt=12, italic=True, color=COLOR_SECONDARY)

# References Callout Box
tbl_ref = doc.add_table(rows=1, cols=1)
tbl_ref.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_ref.autofit = False
tbl_ref.columns[0].width = Inches(6.27)
cell_ref = tbl_ref.rows[0].cells[0]
set_cell_background(cell_ref, HEX_BG_LIGHT)
set_cell_margins(cell_ref, top=140, bottom=140, left=200, right=200)
# Left thick border logic
tcPr = cell_ref._tc.get_or_add_tcPr()
tcBorders = parse_xml(f'''
    <w:tcBorders {nsdecls("w")}>
        <w:top w:val="none"/>
        <w:left w:val="single" w:sz="24" w:space="0" w:color="2B6CB0"/>
        <w:bottom w:val="none"/>
        <w:right w:val="none"/>
    </w:tcBorders>
''')
tcPr.append(tcBorders)

p_ref_head = cell_ref.paragraphs[0]
p_ref_head.paragraph_format.space_after = Pt(4)
run_ref_head = p_ref_head.add_run("Referensi Standar:")
apply_run_font(run_ref_head, size_pt=10, bold=True, color=COLOR_PRIMARY)

p_ref1 = cell_ref.add_paragraph()
p_ref1.paragraph_format.space_after = Pt(2)
p_ref1.paragraph_format.left_indent = Inches(0.2)
run_ref1 = p_ref1.add_run("1. Patterson, D. A., & Hennessy, J. L. (2018). Computer Organization and Design: The Hardware/Software Interface. Morgan Kaufmann.")
apply_run_font(run_ref1, size_pt=10, italic=True)

p_ref2 = cell_ref.add_paragraph()
p_ref2.paragraph_format.space_after = Pt(0)
p_ref2.paragraph_format.left_indent = Inches(0.2)
run_ref2 = p_ref2.add_run("2. Stallings, W. (2016). Computer Organization and Architecture: Designing for Performance (10th Edition). Pearson.")
apply_run_font(run_ref2, size_pt=10, italic=True)

# Spacing after table
doc.add_paragraph().paragraph_format.space_before = Pt(12)

# Helper function for Headings
def add_custom_heading(text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.keep_with_next = True
    if level == 1:
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(6)
        run = p.add_run(text)
        apply_run_font(run, size_pt=14, bold=True, color=COLOR_PRIMARY)
        # Add bottom border accent via XML if possible, or simple underline style
    elif level == 2:
        p.paragraph_format.space_before = Pt(14)
        p.paragraph_format.space_after = Pt(4)
        run = p.add_run(text)
        apply_run_font(run, size_pt=12, bold=True, color=COLOR_SECONDARY)
    return p

# Section 1
add_custom_heading("1. Pengantar Pipelining & Konsep Hazard", level=1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.15
run = p.add_run("Pipelining adalah teknik implementasi di mana beberapa instruksi dieksekusi secara tumpang tindih (overlapped). Analoginya mirip dengan lini perakitan di pabrik modern. Menurut Stallings (2016), pipelining bertujuan untuk meningkatkan throughput CPU, yaitu jumlah instruksi yang dapat diselesaikan per satuan waktu, bukan mempercepat waktu eksekusi satu instruksi tunggal secara individu.")
apply_run_font(run)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
run = p.add_run("Secara umum, arsitektur MIPS/RISC membagi tahapan eksekusi instruksi menjadi 5 tahapan utama (5-Stage Pipeline):")
apply_run_font(run)

stages = [
    ("IF (Instruction Fetch)", "Mengambil instruksi dari memori utama/cache berdasarkan nilai PC (Program Counter)."),
    ("ID (Instruction Decode)", "Menerjemahkan kode instruksi dan membaca data dari register file."),
    ("EX (Execute)", "Melakukan operasi aritmatika/logika melalui ALU atau menghitung alamat memori."),
    ("MEM (Memory Access)", "Membaca data dari atau menulis data ke memori data (jika diperlukan oleh instruksi load/store)."),
    ("WB (Write Back)", "Menuliskan kembali hasil operasi dari tahapan sebelumnya ke dalam register file.")
]

for stg, desc in stages:
    p_item = doc.add_paragraph(style='List Bullet')
    p_item.paragraph_format.space_after = Pt(3)
    p_item.paragraph_format.line_spacing = 1.15
    r_stg = p_item.add_run(f"{stg}: ")
    apply_run_font(r_stg, bold=True, color=COLOR_SECONDARY)
    r_desc = p_item.add_run(desc)
    apply_run_font(r_desc)

# Hazard Definition Box
add_custom_heading("Definisi Hazard", level=2)
tbl_haz = doc.add_table(rows=1, cols=1)
tbl_haz.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_haz.autofit = False
tbl_haz.columns[0].width = Inches(6.27)
cell_haz = tbl_haz.rows[0].cells[0]
set_cell_background(cell_haz, "F8FAFC")
set_cell_margins(cell_haz, top=100, bottom=100, left=150, right=150)
cell_haz._tc.get_or_add_tcPr().append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:color="CBD5E0"/><w:left w:val="single" w:sz="4" w:color="CBD5E0"/><w:bottom w:val="single" w:sz="4" w:color="CBD5E0"/><w:right w:val="single" w:sz="4" w:color="CBD5E0"/></w:tcBorders>'))

p_haz = cell_haz.paragraphs[0]
p_haz.paragraph_format.space_after = Pt(0)
p_haz.paragraph_format.line_spacing = 1.15
r_haz = p_haz.add_run("Hazard adalah situasi di mana pipeline tidak dapat mengeksekusi instruksi selanjutnya pada siklus clock berikutnya karena ada instruksi sebelumnya yang belum selesai dieksekusi atau kondisi data/arsitektur belum terpenuhi. Hazard menyebabkan penurunan performa dan memaksa pipeline untuk melakukan stall (berhenti sejenak dengan memasukkan operasi kosong atau bubble).")
apply_run_font(r_haz, italic=True)

doc.add_paragraph().paragraph_format.space_before = Pt(6)

# Section 2
add_custom_heading("2. Klasifikasi Hazard (3 Jenis Utama)", level=1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
r = p.add_run("Patterson & Hennessy membagi kendala pipeline ke dalam tiga kategori utama yang membutuhkan pendekatan penanganan berbeda baik secara struktural maupun logika logika logika kontrol:")
apply_run_font(r)

# Heading A
add_custom_heading("A. Structural Hazard (Hazard Struktural)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Terjadi ketika hardware atau resource yang sama dibutuhkan oleh dua atau lebih instruksi yang sedang berjalan bersamaan dalam pipeline pada siklus clock yang sama.")
apply_run_font(r)

p_bullet1 = doc.add_paragraph(style='List Bullet')
p_bullet1.paragraph_format.space_after = Pt(2)
r_b1_b = p_bullet1.add_run("Penyebab: ")
apply_run_font(r_b1_b, bold=True)
r_b1_t = p_bullet1.add_run("Keterbatasan alokasi perangkat keras. Sebagai contoh, jika CPU hanya memiliki satu unit port memori tunggal yang digunakan bersama-sama untuk mengambil instruksi (tahap IF) sekaligus membaca/menulis data (tahap MEM).")
apply_run_font(r_b1_t)

p_bullet2 = doc.add_paragraph(style='List Bullet')
p_bullet2.paragraph_format.space_after = Pt(2)
r_b2_b = p_bullet2.add_run("Contoh Kasus: ")
apply_run_font(r_b2_b, bold=True)
r_b2_t = p_bullet2.add_run("Ketika Instruksi 1 berada di tahap MEM (mengakses memori data), secara bersamaan Instruksi 4 berada di tahap IF (mengambil instruksi dari memori). Jika jalurnya tunggal, salah satu harus mengalah.")
apply_run_font(r_b2_t)

p_bullet3 = doc.add_paragraph(style='List Bullet')
p_bullet3.paragraph_format.space_after = Pt(4)
r_b3_b = p_bullet3.add_run("Solusi Teknikal: ")
apply_run_font(r_b3_b, bold=True)
r_b3_t = p_bullet3.add_run("Menerapkan Arsitektur Harvard (memisahkan memori fisik cache menjadi Instruction Cache [I-Cache] dan Data Cache [D-Cache]) atau melakukan penundaan (stalling) secara hardware.")
apply_run_font(r_b3_t)

# Heading B
add_custom_heading("B. Data Hazard (Hazard Data)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Terjadi ketika sebuah instruksi bergantung pada hasil eksekusi instruksi sebelumnya yang masih berada di dalam pipeline dan belum sempat ditulis kembali (Write Back) ke register tujuan.")
apply_run_font(r)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("William Stallings mengklasifikasikan 3 kategori ketergantungan data:")
apply_run_font(r)

data_deps = [
    ("RAW (Read After Write)", "True Dependency. Instruksi B mencoba membaca register sebelum Instruksi A selesai menulis data baru tersebut. Ini adalah jenis hazard data yang paling umum terjadi pada pipeline standard."),
    ("WAR (Write After Read)", "Anti-dependency. Instruksi B mencoba menulis ke register tujuan sebelum Instruksi A sempat membacanya. Biasanya terjadi pada arsitektur canggih Out-of-Order execution."),
    ("WAW (Write After Write)", "Output Dependency. Instruksi B mencoba menulis ke register tujuan sebelum Instruksi A selesai melakukan penulisan. Hal ini dapat mengacaukan urutan final data valid.")
]
for title, text in data_deps:
    p_dep = doc.add_paragraph(style='List Bullet')
    p_dep.paragraph_format.space_after = Pt(2)
    r_t = p_dep.add_run(f"{title} – ")
    apply_run_font(r_t, bold=True)
    r_txt = p_dep.add_run(text)
    apply_run_font(r_txt)

# Assembly Code block style
p_code_title = doc.add_paragraph()
p_code_title.paragraph_format.space_before = Pt(6)
p_code_title.paragraph_format.space_after = Pt(2)
r_ct = p_code_title.add_run("Contoh Kasus RAW pada Kode Assembly:")
apply_run_font(r_ct, size_pt=10, bold=True, color=COLOR_SECONDARY)

tbl_code = doc.add_table(rows=1, cols=1)
tbl_code.alignment = WD_TABLE_ALIGNMENT.CENTER
tbl_code.columns[0].width = Inches(6.27)
c_code = tbl_code.rows[0].cells[0]
set_cell_background(c_code, "F1F5F9")
set_cell_margins(c_code, top=80, bottom=80, left=120, right=120)
c_code._tc.get_or_add_tcPr().append(parse_xml(f'<w:tcBorders {nsdecls("w")}><w:top w:val="single" w:sz="4" w:color="CBD5E0"/><w:left w:val="single" w:sz="4" w:color="CBD5E0"/><w:bottom w:val="single" w:sz="4" w:color="CBD5E0"/><w:right w:val="single" w:sz="4" w:color="CBD5E0"/></w:tcBorders>'))

p_c1 = c_code.paragraphs[0]
p_c1.paragraph_format.space_after = Pt(2)
r_c1 = p_c1.add_run("ADD R1, R2, R3   ; Hasil penjumlahan R1 baru matang setelah tahap EX/MEM")
apply_run_font(r_c1, name="Consolas", size_pt=10, color=RGBColor(21, 128, 61))

p_c2 = c_code.add_paragraph()
p_c2.paragraph_format.space_after = Pt(0)
r_c2 = p_c2.add_run("SUB R4, R1, R5   ; Membutuhkan R1 di tahap ID, padahal ADD belum melakukan Write Back")
apply_run_font(r_c2, name="Consolas", size_pt=10, color=RGBColor(185, 28, 28))

# Solutions for Data Hazard
p_sol_title = doc.add_paragraph()
p_sol_title.paragraph_format.space_before = Pt(6)
p_sol_title.paragraph_format.space_after = Pt(2)
r_st = p_sol_title.add_run("Solusi Penanganan Data Hazard:")
apply_run_font(r_st, bold=True)

data_sols = [
    ("Forwarding / Bypassing: ", "Menyediakan jalur hardware internal tambahan untuk menyalurkan hasil bypass langsung dari output ALU (tahap EX) menuju input ALU instruksi berikutnya tanpa menunggu proses tulis ke register file selesai."),
    ("Stall / Interlock Bubble: ", "Jika data berasal dari operasi pembacaan memori (seperti instruksi Load Word [LW]), hardware terpaksa menyisipkan bubble (1 clock cycle) karena data baru tersedia setelah tahap MEM."),
    ("Compiler Code Reordering: ", "Kompiler secara pintar mengatur ulang urutan instruksi independen yang tidak saling bergantung untuk mengisi celah kekosongan waktu agar eksekusi berjalan linear tanpa stall.")
]
for s_title, s_text in data_sols:
    p_s = doc.add_paragraph(style='List Bullet')
    p_s.paragraph_format.space_after = Pt(2)
    r_st = p_s.add_run(s_title)
    apply_run_font(r_st, bold=True, color=COLOR_SECONDARY)
    r_txt = p_s.add_run(s_text)
    apply_run_font(r_txt)

# Heading C
add_custom_heading("C. Control Hazard / Branch Hazard (Hazard Kontrol)", level=2)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(4)
r = p.add_run("Terjadi ketika pipeline harus mengambil keputusan percabangan berdasarkan hasil instruksi kontrol (seperti Branch Equal [BEQ] atau Jump [J]), tetapi alamat instruksi berikutnya belum divalidasi karena kalkulasi kondisi branch belum selesai dilakukan.")
apply_run_font(r)

p_sol_ctrl = doc.add_paragraph()
p_sol_ctrl.paragraph_format.space_after = Pt(2)
r_sct = p_sol_ctrl.add_run("Metode Penanganan Control Hazard:")
apply_run_font(r_sct, bold=True)

ctrl_sols = [
    ("Stall on Branch", "Menunda proses fetch instruksi berikutnya secara total sampai kalkulasi target percabangan terkonfirmasi secara absolut."),
    ("Branch Prediction (Statik/Dinamik)", "Memprediksi arah percabangan lebih awal. Jika tebakan salah, instruksi yang terlanjur masuk dibersihkan (pipeline flush/clear). Dinamik menggunakan Branch History Table (BHT)."),
    ("Delayed Branch", "Teknik berbasis software/arsitektur RISC lama di mana instruksi yang disisipkan tepat di satu slot setelah branch (delay slot) akan selalu dijalankan secara mutlak baik branch sukses maupun gagal.")
]
for cs_title, cs_text in ctrl_sols:
    p_cs = doc.add_paragraph(style='List Bullet')
    p_cs.paragraph_format.space_after = Pt(2)
    r_cst = p_cs.add_run(f"{cs_title}: ")
    apply_run_font(r_cst, bold=True, color=COLOR_SECONDARY)
    r_ctxt = p_cs.add_run(cs_text)
    apply_run_font(r_ctxt)

# Section 3: Summary Table
add_custom_heading("3. Ringkasan Solusi Penanganan Hazard", level=1)

# Table configuration
table = doc.add_table(rows=4, cols=4)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
set_table_borders(table, color="CBD5E0")

# Adjust column widths
col_widths = [Inches(1.2), Inches(1.8), Inches(1.6), Inches(1.67)]
for row in table.rows:
    for i, w in enumerate(col_widths):
        row.cells[i].width = w

headers = ["Jenis Hazard", "Dampak Pipeline", "Solusi Hardware", "Solusi Software / Compiler"]
hdr_row = table.rows[0]
# Set header repeatable
trPr = hdr_row._tr.get_or_add_trPr()
trPr.append(OxmlElement('w:tblHeader'))

for i, text in enumerate(headers):
    cell = hdr_row.cells[i]
    set_cell_background(cell, "1A365D")
    set_cell_margins(cell, top=120, bottom=120, left=100, right=100)
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(0)
    r = p.add_run(text)
    apply_run_font(r, size_pt=10, bold=True, color=RGBColor(255, 255, 255))

data_rows = [
    ["Structural", "Perebutan resource / gerbang hardware secara simultan.", "Pemisahan fisik memori (I-Cache & D-Cache), replikasi fungsional.", "-"],
    ["Data (RAW)", "Pembacaan operand data register yang belum siap.", "Forwarding/Bypassing, Pipeline Interlock Hardware.", "Instruction Scheduling (Mengubah urutan instruksi aman)."],
    ["Control", "Ketidakpastian alur arah eksekusi percabangan.", "Branch Prediction (Statik/Dinamik), Spekulasi, Flush Pipeline.", "Delayed Branch (Mengisi slot kosong dengan instruksi berguna)."]
]

for row_idx, data in enumerate(data_rows):
    row = table.rows[row_idx + 1]
    # Zebra striping
    bg_color = HEX_BG_LIGHT if row_idx % 2 == 1 else "FFFFFF"
    for col_idx, text in enumerate(data):
        cell = row.cells[col_idx]
        set_cell_background(cell, bg_color)
        set_cell_margins(cell, top=100, bottom=100, left=100, right=100)
        p = cell.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1.15
        r = p.add_run(text)
        apply_run_font(r, size_pt=10)

doc.add_paragraph().paragraph_format.space_before = Pt(12)

# Section 4: Conclusion
add_custom_heading("4. Kesimpulan & Metrik Performa", level=1)
p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(6)
p.paragraph_format.line_spacing = 1.15
r = p.add_run("Berdasarkan analisis Patterson & Hennessy, nilai CPI (Clock Cycles Per Instruction) yang ideal pada sistem arsitektur pipeline bernilai sempurna 1. Namun, akibat kemunculan hazard yang memaksa terjadinya stall atau penundaan siklus, performa aktual dihitung berdasarkan formula matematika berikut:")
apply_run_font(r)

# Formula presentation
p_form = doc.add_paragraph()
p_form.paragraph_format.space_before = Pt(8)
p_form.paragraph_format.space_after = Pt(8)
p_form.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_form = p_form.add_run("CPI Aktual = CPI Ideal (1) + Siklus Stall Pipeline")
apply_run_font(r_form, name="Times New Roman", size_pt=12, bold=True, color=COLOR_PRIMARY)

p = doc.add_paragraph()
p.paragraph_format.space_after = Pt(0)
p.paragraph_format.line_spacing = 1.15
r = p.add_run("Oleh karena itu, fokus utama dari desain arsitektur komputer modern saat ini adalah merancang mekanisme mitigasi guna meminimalkan kontribusi nilai 'Siklus Stall Pipeline'. Caranya adalah dengan mengombinasikan optimasi hardware (forwarding agresif dan dynamic branch prediction) serta optimasi perangkat lunak via cerdasnya proses kompilasi kode pemrograman.")
apply_run_font(r)

# Save document
doc.save("Materi_Arkom_Pipeline_Instruction_Hazard.docx")
print("Successfully generated Word Document.")